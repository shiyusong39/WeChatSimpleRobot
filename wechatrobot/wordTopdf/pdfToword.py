# -*- coding: UTF-8 -*-
'''
PDF 转 Word
需要的包
#pip install pdfminer3k
#pip install python-docx
'''
import os


from io import StringIO
from io import open

from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document

#读取pdf文件
def read_pdf(pdfPath):
    with open(pdfPath,'rb') as pdf_file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, pdf_file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content

#移除特殊字符（制表符、转义符、换行符等）
def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)

#生成word文件
def creat_word(content,wordPath):
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(
            remove_control_characters(line), style=None
        )
        # paragraph = doc.paragraphs()
        # paragraph.add_run(remove_control_characters(line))
    doc.save(wordPath)

def pdf_to_word(pdfPath,wordPath):
    #读取pdf文件
    content = read_pdf(pdfPath)
    if content:
        #生成word
        creat_word(content,wordPath)
        print('转化完成！')
    else:
        print('读取PDF文件失败！')

if __name__ == '__main__':
    print('正在转化。。。')
    pdfPath='D://20200426.pdf'
    wordPath='D://tst.docx'
    pdf_to_word(pdfPath,wordPath)